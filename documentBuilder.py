from docx import Document
import random

class StringTable:
    def __init__(self):
        self.strings = [""]
        self.current_index = 0
        self.move_index = 0

    def insert_to_current(self, text):
        self.strings[self.current_index] = self.strings[self.current_index][:self.move_index] + text + self.strings[ self.current_index][self.move_index:]
        self.move_index += len(text)

    def delete_current(self):
        del self.strings[self.current_index]
        if self.current_index >= len(self.strings):
            self.current_index = len(self.strings) - 1
        self.move_index = len(self.strings[self.current_index])

    def print_current(self):
        return self.strings[self.current_index]

    def go_up(self):
        if self.current_index > 0:
            self.current_index -= 1
            self.move_index = len(self.strings[self.current_index])
            return self.strings[self.current_index]
        else:
            return None

    def go_down(self):
        if self.current_index < len(self.strings) - 1:
            self.current_index += 1
            self.move_index = len(self.strings[self.current_index])
            return self.strings[self.current_index]
        else:
            return None

    def enter_string(self):
        self.strings.insert(self.current_index + 1, "")
        self.current_index += 1
        self.move_index = len(self.strings[self.current_index])

    def print_all(self):
        for string in self.strings:
            print(string)
    def print_current_strict(self):
        tmp=""
        for i in range(self.move_index-1):
            tmp+=" "
        tmp+="|"
        print(tmp)
        print(self.strings[self.current_index])

    def go_left(self):
        if(self.move_index!=0):
            self.move_index=self.move_index-1

    def go_right(self):
        if(self.move_index<len(self.strings[self.current_index])):
            self.move_index=self.move_index+1

    def get_move(self):
        return self.move_index

    def delete_in_move(self):
        if self.move_index > 0:
            self.strings[self.current_index] = self.strings[self.current_index][:self.move_index - 1] + self.strings[self.current_index][self.move_index:]
            self.move_index -= 1

    def delete_All(self):
        self.strings = [""]
        self.current_index = 0
        self.move_index = 0

    def get_title_from_text(self):
        headings = []
        for tmp_string in self.strings:
            if tmp_string.startswith("!h"):
                headings.append(tmp_string.split(" ", 1)[1])
        if headings:
            return headings[0]
        elif any(self.strings):
            return random.choice(" ".join(self.strings).split())
        else:
            return "blank"

    def save_to_document(self):
        doc = Document()
        title = self.get_title_from_text()
        for tmp_string in self.strings:
            if tmp_string.startswith("!"):
                mode, content = tmp_string.split(" ", 1)
                mode = mode[1:]

                if mode.startswith("h"):
                    level = int(mode[1])
                    doc.add_heading(content, level=level)
                elif mode == "ln":
                    doc.add_paragraph(content, style="List Number")
                elif mode == "lk":
                    doc.add_paragraph(content, style="List Bullet")
                elif mode == "bo":
                    p = doc.add_paragraph()
                    p.add_run(content).bold = True
                elif mode == "it":
                    p = doc.add_paragraph()
                    p.add_run(content).italic = True
                elif mode == "pd":
                    p = doc.add_paragraph()
                    p.add_run(content).underline = True
            else:
                doc.add_paragraph(tmp_string)

            doc.save(title+".docx")